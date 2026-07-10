from __future__ import annotations

from io import BytesIO

from docx import Document
from pypdf import PdfWriter

from app.services.knowledge_parser import ScannedPdfError, parse_document, split_text


def test_txt_parser_splits_long_narrative_within_configured_limit() -> None:
    source = ("灵山大佛是景区文化核心。" * 120).encode("utf-8")
    chunks = parse_document(source, "guide.txt", "text/plain")
    assert len(chunks) > 1
    assert all(0 < len(chunk.content) <= 800 for chunk in chunks)
    assert all(chunk.source_locator == "guide.txt" for chunk in chunks)


def test_docx_table_preserves_spot_identity() -> None:
    document = Document()
    document.add_heading("景点资料", level=1)
    table = document.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "景点ID"
    table.rows[0].cells[1].text = "景点名称"
    table.rows[0].cells[2].text = "文化亮点"
    table.rows[1].cells[0].text = "LS-001"
    table.rows[1].cells[1].text = "灵山大佛"
    table.rows[1].cells[2].text = "展示佛教文化与艺术"
    buffer = BytesIO()
    document.save(buffer)

    chunks = parse_document(buffer.getvalue(), "spots.docx")
    structured = next(chunk for chunk in chunks if chunk.spot_id == "LS-001")
    assert structured.spot_name == "灵山大佛"
    assert "文化亮点" in structured.content


def test_scanned_pdf_is_rejected_with_ocr_guidance() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buffer = BytesIO()
    writer.write(buffer)
    try:
        parse_document(buffer.getvalue(), "scan.pdf", "application/pdf")
    except ScannedPdfError as exc:
        assert "OCR" in str(exc)
    else:
        raise AssertionError("blank PDF should be rejected as scanned")


def test_chunk_overlap_keeps_context_between_adjacent_chunks() -> None:
    chunks = split_text("甲" * 750 + "。" + "乙" * 750)
    assert len(chunks) >= 2
    assert chunks[0][-120:] == chunks[1][:120]
