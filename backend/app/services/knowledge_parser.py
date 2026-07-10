from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
import re

from docx import Document
from pypdf import PdfReader


MAX_CHUNK_CHARS = 800
CHUNK_OVERLAP_CHARS = 120


class UnsupportedDocumentError(ValueError):
    pass


class ScannedPdfError(UnsupportedDocumentError):
    pass


@dataclass(frozen=True)
class ParsedChunk:
    content: str
    section: str | None = None
    spot_id: str | None = None
    spot_name: str | None = None
    source_locator: str | None = None


def _clean(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def split_text(text: str, max_chars: int = MAX_CHUNK_CHARS, overlap: int = CHUNK_OVERLAP_CHARS) -> list[str]:
    text = _clean(text)
    if not text:
        return []
    if len(text) <= max_chars:
        return [text]

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + max_chars, len(text))
        if end < len(text):
            boundary = max(text.rfind(mark, start + max_chars // 2, end) for mark in "。！？；\n")
            if boundary >= start + max_chars // 2:
                end = boundary + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(end - overlap, start + 1)
    return chunks


def _narrative_chunks(paragraphs: list[str], source_name: str) -> list[ParsedChunk]:
    chunks: list[ParsedChunk] = []
    section: str | None = None
    buffer: list[str] = []

    def flush() -> None:
        nonlocal buffer
        if not buffer:
            return
        for part in split_text("\n".join(buffer)):
            chunks.append(ParsedChunk(content=part, section=section, source_locator=source_name))
        buffer = []

    for paragraph in paragraphs:
        value = _clean(paragraph)
        if not value:
            continue
        if len(value) <= 40 and (value.endswith("：") or re.match(r"^(第?[一二三四五六七八九十0-9]+[、.．]|[一二三四五六七八九十]+、)", value)):
            flush()
            section = value.rstrip("：")
        else:
            buffer.append(value)
    flush()
    return chunks


def _table_chunks(doc: Document, source_name: str) -> list[ParsedChunk]:
    result: list[ParsedChunk] = []
    for table_index, table in enumerate(doc.tables, start=1):
        rows = [[_clean(cell.text) for cell in row.cells] for row in table.rows]
        rows = [row for row in rows if any(row)]
        if len(rows) < 2:
            continue
        headers = rows[0]
        normalized = [re.sub(r"\s", "", header).lower() for header in headers]
        spot_id_index = next((i for i, value in enumerate(normalized) if "景点id" in value or value in {"id", "景点编号"}), None)
        spot_name_index = next((i for i, value in enumerate(normalized) if "景点名称" in value or value in {"名称", "景点"}), None)
        for row_index, row in enumerate(rows[1:], start=2):
            values = [f"{headers[index]}：{value}" for index, value in enumerate(row) if index < len(headers) and value]
            if not values:
                continue
            spot_id = row[spot_id_index] if spot_id_index is not None and spot_id_index < len(row) else None
            spot_name = row[spot_name_index] if spot_name_index is not None and spot_name_index < len(row) else None
            content = "；".join(values)
            for part in split_text(content):
                result.append(
                    ParsedChunk(
                        content=part,
                        section="结构化景点资料",
                        spot_id=spot_id or None,
                        spot_name=spot_name or None,
                        source_locator=f"{source_name} / 表{table_index}第{row_index}行",
                    )
                )
    return result


def parse_txt(data: bytes, source_name: str) -> list[ParsedChunk]:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            text = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise UnsupportedDocumentError("TXT 文件编码无法识别，请保存为 UTF-8 后重新上传")
    return _narrative_chunks(text.splitlines(), source_name)


def parse_docx(data: bytes, source_name: str) -> list[ParsedChunk]:
    try:
        doc = Document(BytesIO(data))
    except Exception as exc:  # python-docx reports different ZIP/XML failures by version
        raise UnsupportedDocumentError("DOCX 文件无法解析") from exc
    chunks = _narrative_chunks([paragraph.text for paragraph in doc.paragraphs], source_name)
    chunks.extend(_table_chunks(doc, source_name))
    return chunks


def parse_pdf(data: bytes, source_name: str) -> list[ParsedChunk]:
    try:
        reader = PdfReader(BytesIO(data))
        pages = [_clean(page.extract_text() or "") for page in reader.pages]
    except Exception as exc:
        raise UnsupportedDocumentError("PDF 文件无法读取") from exc
    text = "\n".join(page for page in pages if page)
    if not text:
        raise ScannedPdfError("扫描型 PDF，需 OCR 后重新上传")
    return _narrative_chunks(text.splitlines(), source_name)


def parse_document(data: bytes, filename: str, mime_type: str | None = None) -> list[ParsedChunk]:
    suffix = Path(filename).suffix.lower()
    if suffix == ".txt" or mime_type == "text/plain":
        chunks = parse_txt(data, filename)
    elif suffix == ".docx" or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        chunks = parse_docx(data, filename)
    elif suffix == ".pdf" or mime_type == "application/pdf":
        chunks = parse_pdf(data, filename)
    else:
        raise UnsupportedDocumentError("仅支持 DOCX、TXT 和可提取文本的 PDF 文件")
    if not chunks:
        raise UnsupportedDocumentError("未从文件中提取到可索引文本")
    return chunks
