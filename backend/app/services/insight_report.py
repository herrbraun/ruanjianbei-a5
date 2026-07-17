from __future__ import annotations

import json
import re
from datetime import date, datetime, timedelta, timezone
from io import BytesIO
from zoneinfo import ZoneInfo

import httpx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from pydantic import ValidationError
from sqlalchemy import func, select, update
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import Session

from app.config import settings
from app.crud.insights import get_guide_dashboard
from app.database import SessionLocal
from app.models.guide import InsightReportSchedule, ScenicInsightReport
from app.models.knowledge import ScenicArea
from app.schemas.insights import ReportNarrative


class InsightReportError(RuntimeError):
    pass


MAX_GENERATION_ATTEMPTS = 3
MAX_CATCHUP_REPORTS_PER_TYPE = 7


def report_deduplication_key(scenic_area_id: int, period_type: str, period_start: date, period_end: date) -> str:
    return f"{scenic_area_id}:{period_type}:{period_start.isoformat()}:{period_end.isoformat()}"


def build_report_snapshot(db: Session, scenic_area_id: int, period_start: date, period_end: date) -> dict:
    dashboard = get_guide_dashboard(db, scenic_area_id, period_start, period_end)
    return {
        **dashboard["metrics"],
        "top_topics": dashboard["topic_distribution"],
        "popular_questions": dashboard["popular_questions"],
        "sentiment_trend": dashboard["sentiment_trend"],
        "satisfaction_trend": dashboard["satisfaction_trend"],
    }


def create_or_get_report(
    db: Session,
    *,
    scenic_area_id: int,
    period_type: str,
    period_start: date,
    period_end: date,
    trigger_source: str,
    created_by_user_id: int | None,
) -> tuple[ScenicInsightReport, bool]:
    key = report_deduplication_key(scenic_area_id, period_type, period_start, period_end)
    existing = db.scalar(
        select(ScenicInsightReport)
        .where(
            ScenicInsightReport.scenic_area_id == scenic_area_id,
            ScenicInsightReport.period_type == period_type,
            ScenicInsightReport.period_start == period_start,
            ScenicInsightReport.period_end == period_end,
        )
        .order_by(ScenicInsightReport.id)
    )
    if existing is not None:
        if existing.deduplication_key is None:
            existing.deduplication_key = key
            try:
                db.commit()
            except IntegrityError:
                db.rollback()
            else:
                db.refresh(existing)
        return existing, False
    report = ScenicInsightReport(
        scenic_area_id=scenic_area_id,
        period_type=period_type,
        period_start=period_start,
        period_end=period_end,
        metrics_snapshot=build_report_snapshot(db, scenic_area_id, period_start, period_end),
        generation_status="pending",
        trigger_source=trigger_source,
        deduplication_key=key,
        created_by_user_id=created_by_user_id,
    )
    db.add(report)
    try:
        db.commit()
    except IntegrityError:
        db.rollback()
        existing = db.scalar(select(ScenicInsightReport).where(ScenicInsightReport.deduplication_key == key))
        if existing is None:
            raise
        return existing, False
    db.refresh(report)
    return report, True


def get_or_create_report_schedule(db: Session, scenic_area_id: int) -> InsightReportSchedule:
    schedule = db.scalar(
        select(InsightReportSchedule).where(InsightReportSchedule.scenic_area_id == scenic_area_id)
    )
    if schedule is None:
        schedule = InsightReportSchedule(scenic_area_id=scenic_area_id)
        db.add(schedule)
        try:
            db.commit()
        except IntegrityError:
            db.rollback()
            schedule = db.scalar(
                select(InsightReportSchedule).where(
                    InsightReportSchedule.scenic_area_id == scenic_area_id
                )
            )
            if schedule is None:
                raise
        else:
            db.refresh(schedule)
    return schedule


def ensure_default_report_schedules(db: Session) -> int:
    existing_area_ids = set(db.scalars(select(InsightReportSchedule.scenic_area_id)))
    scenic_area_ids = set(
        db.scalars(select(ScenicArea.id).where(ScenicArea.is_enabled.is_(True)))
    )
    missing_area_ids = scenic_area_ids - existing_area_ids
    for scenic_area_id in missing_area_ids:
        db.add(InsightReportSchedule(scenic_area_id=scenic_area_id))
    if missing_area_ids:
        try:
            db.commit()
        except IntegrityError:
            # Another worker may have inserted the same defaults between the
            # read and commit. The unique scenic-area key makes that race safe.
            db.rollback()
            return 0
    return len(missing_area_ids)


def _latest_daily_period(schedule: InsightReportSchedule, now: datetime) -> tuple[date, date]:
    local_now = now.astimezone(ZoneInfo(schedule.timezone))
    scheduled_date = local_now.date()
    if local_now.time().replace(tzinfo=None) < schedule.daily_run_time:
        scheduled_date -= timedelta(days=1)
    report_date = scheduled_date - timedelta(days=1)
    return report_date, report_date


def _latest_weekly_period(schedule: InsightReportSchedule, now: datetime) -> tuple[date, date]:
    local_now = now.astimezone(ZoneInfo(schedule.timezone))
    week_start = local_now.date() - timedelta(days=local_now.weekday())
    scheduled_date = week_start + timedelta(days=schedule.weekly_weekday)
    if scheduled_date > local_now.date() or (
        scheduled_date == local_now.date()
        and local_now.time().replace(tzinfo=None) < schedule.weekly_run_time
    ):
        scheduled_date -= timedelta(days=7)
    period_end = scheduled_date - timedelta(days=1)
    return period_end - timedelta(days=6), period_end


def _missing_scheduled_periods(
    db: Session,
    schedule: InsightReportSchedule,
    period_type: str,
    latest_start: date,
    latest_end: date,
) -> list[tuple[date, date]]:
    first_scheduled_end = db.scalar(
        select(func.min(ScenicInsightReport.period_end)).where(
            ScenicInsightReport.scenic_area_id == schedule.scenic_area_id,
            ScenicInsightReport.period_type == period_type,
            ScenicInsightReport.trigger_source == "scheduled",
            ScenicInsightReport.period_end <= latest_end,
        )
    )
    existing_ends = set(
        db.scalars(
            select(ScenicInsightReport.period_end).where(
                ScenicInsightReport.scenic_area_id == schedule.scenic_area_id,
                ScenicInsightReport.period_type == period_type,
                ScenicInsightReport.period_end <= latest_end,
            )
        )
    )
    if first_scheduled_end is None:
        return [] if latest_end in existing_ends else [(latest_start, latest_end)]

    step_days = 1 if period_type == "daily" else 7
    if (latest_end - first_scheduled_end).days % step_days:
        return [] if latest_end in existing_ends else [(latest_start, latest_end)]

    missing: list[tuple[date, date]] = []
    candidate_end = first_scheduled_end + timedelta(days=step_days)
    while candidate_end <= latest_end and len(missing) < MAX_CATCHUP_REPORTS_PER_TYPE:
        if candidate_end not in existing_ends:
            candidate_start = (
                candidate_end if period_type == "daily" else candidate_end - timedelta(days=6)
            )
            missing.append((candidate_start, candidate_end))
        candidate_end += timedelta(days=step_days)
    return missing


def create_due_reports(db: Session, now: datetime | None = None) -> list[int]:
    current = now or datetime.now(timezone.utc)
    created_ids: list[int] = []
    schedules = list(
        db.scalars(
            select(InsightReportSchedule)
            .join(ScenicArea, ScenicArea.id == InsightReportSchedule.scenic_area_id)
            .where(ScenicArea.is_enabled.is_(True))
            .order_by(InsightReportSchedule.id)
        )
    )
    for schedule in schedules:
        periods: list[tuple[str, date, date]] = []
        if schedule.daily_enabled:
            start, end = _latest_daily_period(schedule, current)
            periods.extend(
                ("daily", period_start, period_end)
                for period_start, period_end in _missing_scheduled_periods(
                    db, schedule, "daily", start, end
                )
            )
        if schedule.weekly_enabled:
            start, end = _latest_weekly_period(schedule, current)
            periods.extend(
                ("weekly", period_start, period_end)
                for period_start, period_end in _missing_scheduled_periods(
                    db, schedule, "weekly", start, end
                )
            )
        for period_type, start, end in periods:
            report, created = create_or_get_report(
                db,
                scenic_area_id=schedule.scenic_area_id,
                period_type=period_type,
                period_start=start,
                period_end=end,
                trigger_source="scheduled",
                created_by_user_id=None,
            )
            if created:
                created_ids.append(report.id)
    return created_ids


def recover_stale_reports(db: Session, stale_before: datetime) -> int:
    rows = list(
        db.scalars(
            select(ScenicInsightReport).where(
                ScenicInsightReport.generation_status == "processing",
                ScenicInsightReport.processing_started_at < stale_before,
            )
        )
    )
    for report in rows:
        report.generation_status = (
            "pending" if report.generation_attempts < MAX_GENERATION_ATTEMPTS else "failed"
        )
        report.error_message = "任务中断，等待重试"
        report.processing_started_at = None
    db.commit()
    return len(rows)


def pending_report_ids(db: Session, limit: int = 20) -> list[int]:
    return list(
        db.scalars(
            select(ScenicInsightReport.id)
            .where(ScenicInsightReport.generation_status == "pending")
            .order_by(ScenicInsightReport.created_at)
            .limit(limit)
        )
    )


def _parse_report(payload: dict) -> ReportNarrative:
    try:
        raw = payload["choices"][0]["message"]["content"]
        cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw.strip(), flags=re.IGNORECASE)
        return ReportNarrative.model_validate(json.loads(cleaned))
    except (KeyError, IndexError, TypeError, json.JSONDecodeError, ValidationError) as exc:
        raise InsightReportError("报告模型返回格式无效") from exc


def generate_insight_report(snapshot: dict, *, client: httpx.Client | None = None) -> ReportNarrative:
    if not settings.dashscope_api_key or settings.dashscope_api_key == "your_dashscope_api_key":
        raise InsightReportError("未配置 DASHSCOPE_API_KEY，无法生成感受度报告")
    own = client is None
    active = client or httpx.Client(timeout=60.0)
    try:
        system = (
            "你是景区运营分析师，只能依据聚合指标输出纯JSON，不得改写数字，禁止遗漏或改名字段。"
            "必须包含summary字符串、attention_points恰好3条、risk_findings至少1条、recommendations 3至5条。"
            "示例结构：{\"summary\":\"总体结论\",\"attention_points\":[\"重点1\",\"重点2\",\"重点3\"],"
            "\"risk_findings\":[\"风险1\"],\"recommendations\":[\"建议1\",\"建议2\",\"建议3\"]}。"
        )
        user = json.dumps(snapshot, ensure_ascii=False)
        for attempt in range(2):
            response = active.post(
                f"{settings.llm_base_url.rstrip('/')}/chat/completions",
                headers={"Authorization": f"Bearer {settings.dashscope_api_key}", "Content-Type": "application/json"},
                json={
                    "model": settings.insight_report_model,
                    "messages": [{"role": "system", "content": system}, {"role": "user", "content": user}],
                    "temperature": 0.2,
                    "stream": False,
                    "response_format": {"type": "json_object"},
                },
            )
            if response.is_error:
                raise InsightReportError(f"报告模型调用失败（HTTP {response.status_code}）")
            try:
                return _parse_report(response.json())
            except InsightReportError:
                if attempt:
                    raise
                raw = response.json().get("choices", [{}])[0].get("message", {}).get("content", "")
                user = f"请严格按既定四个字段修复下面的JSON，补齐列表数量，只输出JSON：\n{raw}"
        raise InsightReportError("报告模型没有返回有效结果")
    except httpx.HTTPError as exc:
        raise InsightReportError("报告模型连接失败") from exc
    finally:
        if own:
            active.close()


def process_report(report_id: int) -> None:
    with SessionLocal() as db:
        claim = db.execute(
            update(ScenicInsightReport)
            .where(
                ScenicInsightReport.id == report_id,
                ScenicInsightReport.generation_status == "pending",
            )
            .values(
                generation_status="processing",
                generation_attempts=ScenicInsightReport.generation_attempts + 1,
                processing_started_at=datetime.now(timezone.utc),
                error_message=None,
            )
            .execution_options(synchronize_session=False)
        )
        db.commit()
        if claim.rowcount != 1:
            return
        report = db.get(ScenicInsightReport, report_id)
        if report is None:
            return
        try:
            has_data = bool(
                report.metrics_snapshot.get("service_visitors")
                or report.metrics_snapshot.get("question_count")
            )
            narrative = (
                generate_insight_report(report.metrics_snapshot)
                if has_data
                else ReportNarrative(
                    summary="本周期暂无有效游客互动数据，暂不能形成趋势判断。",
                    attention_points=["服务人次为零", "暂无有效问答", "暂无满意度评价"],
                    risk_findings=["数据量不足，无法识别稳定风险"],
                    recommendations=["确认导览入口正常开放", "持续收集游客互动", "下一周期继续观察数据"],
                )
            )
            report.summary = narrative.summary
            report.attention_points = narrative.attention_points
            report.risk_findings = narrative.risk_findings
            report.recommendations = narrative.recommendations
            report.generation_model = settings.insight_report_model if has_data else "deterministic-empty-period"
            report.generation_status = "completed"
            report.generated_at = datetime.now(timezone.utc)
            report.processing_started_at = None
        except Exception as exc:
            report.generation_status = (
                "pending" if report.generation_attempts < MAX_GENERATION_ATTEMPTS else "failed"
            )
            report.error_message = str(exc)[:2000]
            report.processing_started_at = None
        db.commit()


def _set_run_font(run: object, font_name: str = "Microsoft YaHei") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def build_report_docx(report: ScenicInsightReport, scenic_name: str) -> bytes:
    if report.generation_status != "completed":
        raise InsightReportError("报告尚未生成完成")
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(22)
    section.right_margin = Mm(22)
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    for style_name in ("Title", "Heading 1", "Heading 2"):
        styles[style_name].font.name = "Microsoft YaHei"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    title = document.add_heading(f"{scenic_name}游客感受度{'日报' if report.period_type == 'daily' else '周报'}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        _set_run_font(run)
        run.font.color.rgb = RGBColor(20, 92, 82)
    subtitle = document.add_paragraph(
        f"统计周期：{report.period_start.isoformat()} 至 {report.period_end.isoformat()}"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        _set_run_font(run)

    metrics = report.metrics_snapshot
    metric_rows = [
        ("服务游客数", str(metrics.get("service_visitors", 0))),
        ("问答数量", str(metrics.get("question_count", 0))),
        ("回答成功率", f"{float(metrics.get('answer_success_rate') or 0) * 100:.1f}%"),
        (
            "平均响应时间",
            "暂无" if metrics.get("average_answer_duration_ms") is None
            else f"{float(metrics['average_answer_duration_ms']) / 1000:.2f} 秒",
        ),
        ("平均满意度", "暂无" if metrics.get("average_rating") is None else f"{metrics['average_rating']} / 5"),
        ("负向反馈率", f"{float(metrics.get('negative_rate') or 0) * 100:.1f}%"),
    ]
    document.add_heading("核心指标", level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "指标"
    table.rows[0].cells[1].text = "结果"
    for label, value in metric_rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run)

    def add_list_section(title_text: str, values: list[str] | None, style: str) -> None:
        document.add_heading(title_text, level=1)
        for value in values or ["暂无"]:
            paragraph = document.add_paragraph(str(value), style=style)
            for run in paragraph.runs:
                _set_run_font(run)

    document.add_heading("总体结论", level=1)
    summary = document.add_paragraph(report.summary or "暂无")
    for run in summary.runs:
        _set_run_font(run)
    add_list_section("重点关注", report.attention_points, "List Bullet")
    add_list_section("风险发现", report.risk_findings, "List Bullet")
    add_list_section("改进建议", report.recommendations, "List Number")

    document.add_heading("热门关注", level=1)
    topics = metrics.get("top_topics") or []
    questions = metrics.get("popular_questions") or []
    for item in topics[:8]:
        paragraph = document.add_paragraph(
            f"{item.get('name', '其他')}：{item.get('count', 0)} 次",
            style="List Bullet",
        )
        for run in paragraph.runs:
            _set_run_font(run)
    if questions:
        document.add_heading("热门问答", level=2)
        for item in questions[:10]:
            paragraph = document.add_paragraph(
                f"{item.get('name', '未命名问题')}（{item.get('count', 0)} 次）",
                style="List Number",
            )
            for run in paragraph.runs:
                _set_run_font(run)

    footer = section.footer.paragraphs[0]
    footer.text = (
        f"生成时间：{report.generated_at.astimezone(ZoneInfo('Asia/Shanghai')).strftime('%Y-%m-%d %H:%M') if report.generated_at else '未知'}"
        f"    分析模型：{report.generation_model or '未知'}"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        _set_run_font(run)
        run.font.size = Pt(8)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
